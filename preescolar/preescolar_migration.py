import os
from docx import Document

from utils import clean_name


def split_into_chunks(arr, chunk_size):
    """
    Splits an array into chunks of a given size.

    :param arr: The input list to be split.
    :param chunk_size: The size of each chunk.
    :return: A list of lists, where each sublist is a chunk of the original list.
    """
    if chunk_size <= 0:
        raise ValueError("Chunk size must be greater than zero.")

    return [arr[i:i + chunk_size] for i in range(0, len(arr), chunk_size)]


def get_preescolar_data(path):
    data = {
        "grado": 0,
        "grupo": '',
        "fecha": ''
    }
    file_name = path.split('\\')[-1].split('_')
    data['fecha'] = file_name[0]
    data['grupo'] = file_name[-1][0]
    return data


def obtain_name(text):
    name = text.lower().replace(':', ' ').replace('alumno', 'estudiante')
    name = name.split('estudiante')

    if len(name) < 2:
        print(name)
    name = name[-1]
    name = clean_name(name)
    if not name:
        print('failed name', text)
    return name


def process_file(path: str) -> dict[str, dict[str, str]]:
    """
    Extracts student names and their corresponding dimension evaluations from a DOCX file.

    The DOCX file is expected to have:
    1. Paragraphs containing student names, formatted as "Nombre del estudiante: [name]".
    2. Tables where the first column contains dimension names and the second column contains descriptions.

    :param path: Path to the DOCX file.
    :return: A dictionary where keys are student names, and values are dictionaries
             containing dimension names as keys and their descriptions as values.
    """
    doc = Document(path)

    # Extract student names by filtering paragraphs that contain "nombre del estudiante"
    students = [obtain_name(para.text.strip()) for para in doc.paragraphs
                if 'nombre' in para.text.lower() and 'estudiante' in para.text.lower().replace('alumno', 'estudiante')]

    # Dictionary of valid dimension names
    replacement_dimensions = {
        'CORPORAL': 'DIMENSION_CORPORAL1',
        'CO RPORAL': 'DIMENSION_CORPORAL1',

        'COGNITIVA': 'DIMENSION_COGNITIVA1',
        'COGNOSCITIVA': 'DIMENSION_COGNITIVA1',
        'COGNOSCITIVO': 'DIMENSION_COGNITIVA1',
        'NOSCITIVO': 'DIMENSION_COGNITIVA1',

        'AFECTIVA': 'DIMENSION_SOCIO_AFECTIVA1',
        'SOCIOAFECTIVA': 'DIMENSION_SOCIO_AFECTIVA1',

        'COMUNICATIVA': 'DIMENSION_COMUNICATIVA1',
        'COGNOSCITIVA COMUNICATIVA': 'DIMENSION_COMUNICATIVA1',

        'ETICA': 'DIMENSION_ETICA1',
        'ETICA Y VALORES': 'DIMENSION_ETICA1',
        'TICA': 'DIMENSION_ETICA1',

        'ESTETICA': 'DIMENSION_ESTETICA1',
        'ARTISRICA': 'DIMENSION_ESTETICA1',
        'ARTISTICA': 'DIMENSION_ESTETICA1',

        'ACTITUDINAL Y VALORATIVA': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL Y VALORES': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL YVALORATIVA': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDIVALORATIVA': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL': 'DIMENSION_ACTITUDINAL1',
        'VALORATIVA': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINA Y': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL Y': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINALY VALORATIVA': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL Y A': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL Y VALORAT': 'DIMENSION_ACTITUDINAL1',
        'VALORATIVA Y ACTITU': 'DIMENSION_ACTITUDINAL1',
        'ACTITUDINAL Y Y VALORATIVA': 'DIMENSION_ACTITUDINAL1',
        'VALORATIVA Y ACTITUDINALÑ': 'DIMENSION_ACTITUDINAL1',
        'VALORATIVA Y ACTITUDINAL': 'DIMENSION_ACTITUDINAL1'
    }
    # Stores tuples of (dimension_name, dimension_description)
    tables = []

    # Extract relevant dimension information from tables
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            # Ensure row has at least two columns
            if cells and len(cells) >= 2:
                dimension_name = clean_name(cells[0].text.strip())
                description = cells[-1].text
                if dimension_name in replacement_dimensions.keys() and description:
                    tables.append((replacement_dimensions[dimension_name], description))

    # Ensure valid mapping between students and dimensions
    if not students or not tables:
        # Return an empty dictionary if either is missing
        return {}

    # Ensure integer division
    dimensions_per_student = int(len(tables) / len(students))
    print(len(tables), len(students), dimensions_per_student)
    if len(students) > 30:
        print('Something might be going here')
    if dimensions_per_student == 0:
        raise ValueError("Mismatch: More students than available dimensions.")

    # Split the extracted dimensions into chunks per student
    student_dimensions = split_into_chunks(tables, dimensions_per_student)
    # print(list(map(lambda x: list(map(lambda y: y[:10], x)), student_dimensions)))
    # Create dictionary mapping students to their corresponding dimensions
    return {student: dict(student_dimensions[i])
            for i, student in enumerate(students)}


def post_to_db(file_path, connection):
    cursor = connection.cursor()
    students_and_dimensions = process_file(file_path)
    group_data = get_preescolar_data(file_path)
    failed_count = 0
    all_queries = []
    for student, dimensions in students_and_dimensions.items():

        print(f"\n{student} from {group_data['grado']} - {group_data['grupo']} {group_data['fecha']}")

        values = (
            student,
            dimensions.get('DIMENSION_COGNITIVA1', None),
            dimensions.get('DIMENSION_COMUNICATIVA1', None),
            dimensions.get('DIMENSION_CORPORAL1', None),
            dimensions.get('DIMENSION_ESTETICA1', None),
            dimensions.get('DIMENSION_ETICA1', None),
            dimensions.get('DIMENSION_SOCIO_AFECTIVA1', None),
            dimensions.get('DIMENSION_ACTITUDINAL1', None),
            group_data['fecha'],
            group_data['grado'],
            group_data['grupo']
        )
        all_queries.append(values)
        for dimension, description in dimensions.items():
            print(f"\t{dimension} {description[:30]}")

            if not description:
                print(f"\thas no description {description[:30]}")
                failed_count += 1
    print(f"\n{len(all_queries)}")
    print(f"\nFailed: {failed_count}")

    # if all_queries:
    #     query = (f"INSERT INTO Calificaciones0 (`Nombre`, `DIMENSION_COGNITIVA1`,`DIMENSION_COMUNICATIVA1`, "
    #              f"`DIMENSION_CORPORAL1`, `DIMENSION_ESTETICA1`, `DIMENSION_ETICA1`, "
    #              f"`DIMENSION_SOCIO_AFECTIVA1`, `DIMENSION_ACTITUDINAL1`, "
    #              f"`Fecha`, `Grado`, `Grupo`) "
    #              f"VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);")
    #     try:
    #         # Execute queries in bulk
    #         cursor.executemany(query, all_queries)
    #         connection.commit()
    #         print(f"{len(all_queries)} records inserted successfully.")
    #     except Exception as e:
    #         connection.rollback()
    #         print(f"Error during SQL execution: {e}")


def process_folder_0(folder_path, connection):
    """Process all SQL files in a folder and save their processed data in the db.
    """

    if not os.path.exists(folder_path):
        print(f"Folder does not exist: {folder_path}")
        return

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)

        if os.path.isfile(file_path) and file_name.lower().endswith('.docx'):
            print(f"\n{'-' * 100}")
            print(f"{file_path}\n")
            post_to_db(file_path, connection)
            print(f"{file_path}\n")
